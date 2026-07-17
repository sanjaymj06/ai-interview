import os
import re
from typing import Any, Dict, List, Optional, Tuple

import fitz  # PyMuPDF
import pdfplumber
from docx import Document

from app.utils.constants import REGEX_PATTERNS, RESUME_SECTIONS


def extract_text_from_pdf(file_path: str) -> Tuple[str, int]:
    text = ""
    page_count = 0
    try:
        doc = fitz.open(file_path)
        page_count = len(doc)
        for page in doc:
            text += page.get_text() + "\n"
        doc.close()
    except Exception:
        text = ""
        page_count = 0
    if not text.strip():
        try:
            with pdfplumber.open(file_path) as pdf:
                page_count = len(pdf.pages)
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception:
            pass
    return text.strip(), page_count


def extract_text_from_docx(file_path: str) -> Tuple[str, int]:
    try:
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        text = "\n".join(paragraphs)
        return text.strip(), 1
    except Exception as e:
        raise ValueError(f"Failed to extract text from DOCX: {str(e)}")


def extract_text(file_path: str) -> Tuple[str, int]:
    ext = file_path.rsplit(".", 1)[-1].lower()
    if ext == "pdf":
        return extract_text_from_pdf(file_path)
    elif ext == "docx":
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def clean_text(text: str) -> str:
    text = re.sub(r"\x0c", "\n", text)
    text = re.sub(r"\t", " ", text)
    text = re.sub(r" +", " ", text)
    lines = text.split("\n")
    cleaned = [line.strip() for line in lines if line.strip()]
    return "\n".join(cleaned)


def extract_email(text: str) -> Optional[str]:
    match = re.search(REGEX_PATTERNS["email"], text)
    return match.group(0) if match else None


def extract_phone(text: str) -> Optional[str]:
    match = re.search(REGEX_PATTERNS["phone"], text)
    return match.group(0) if match else None


def extract_linkedin(text: str) -> Optional[str]:
    match = re.search(REGEX_PATTERNS["linkedin"], text)
    return match.group(0) if match else None


def extract_github(text: str) -> Optional[str]:
    match = re.search(REGEX_PATTERNS["github"], text)
    return match.group(0) if match else None


def extract_urls(text: str) -> List[str]:
    return re.findall(REGEX_PATTERNS["url"], text)


def extract_portfolio_urls(text: str, linkedin: str = None, github: str = None) -> Optional[str]:
    urls = extract_urls(text)
    portfolio = None
    for url in urls:
        if linkedin and url in linkedin:
            continue
        if github and url in github:
            continue
        if not re.search(r"linkedin\.com|github\.com", url):
            portfolio = url
            break
    return portfolio


def extract_name_from_text(text: str) -> Optional[str]:
    try:
        import spacy
        nlp = spacy.load("en_core_web_sm")
        doc = nlp(text[:1000])
        for ent in doc.ents:
            if ent.label_ == "PERSON":
                name = ent.text.strip()
                if 2 <= len(name.split()) <= 4:
                    return name
    except Exception:
        pass
    lines = text.strip().split("\n")
    for line in lines[:5]:
        line = line.strip()
        if not line:
            continue
        if re.search(REGEX_PATTERNS["email"], line):
            continue
        if re.search(REGEX_PATTERNS["phone"], line):
            continue
        if re.search(r"[|@#$%^&*]", line):
            continue
        words = line.split()
        if 2 <= len(words) <= 4 and all(w[0].isupper() for w in words if w):
            return line
    return None


def find_section_boundaries(text: str) -> Dict[str, Tuple[int, int]]:
    lines = text.split("\n")
    sections: Dict[str, Tuple[int, int]] = {}
    section_positions: List[Tuple[int, str]] = []
    for i, line in enumerate(lines):
        cleaned = line.strip().lower().rstrip(":")
        for section_name, headers in RESUME_SECTIONS.items():
            if cleaned in headers or cleaned.rstrip(":") in headers:
                section_positions.append((i, section_name))
                break
    section_positions.sort(key=lambda x: x[0])
    for idx, (pos, name) in enumerate(section_positions):
        start = pos
        if idx + 1 < len(section_positions):
            end = section_positions[idx + 1][0]
        else:
            end = len(lines)
        sections[name] = (start, end)
    return sections


def extract_skills_from_text(text: str) -> List[str]:
    from app.utils.constants import SKILL_CATEGORIES
    found_skills = []
    text_lower = text.lower()
    for category, skills in SKILL_CATEGORIES.items():
        for skill in skills:
            pattern = r"\b" + re.escape(skill) + r"\b"
            if re.search(pattern, text_lower):
                if skill not in found_skills:
                    found_skills.append(skill)
    return found_skills


def extract_education(text: str, sections: Dict[str, Tuple[int, int]]) -> List[Dict[str, Any]]:
    education = []
    lines = text.split("\n")
    if "education" in sections:
        start, end = sections["education"]
        edu_text = "\n".join(lines[start:end])
        degree_patterns = [
            r"(Bachelor[\s]*(?:of|in|Science|Arts|Engineering|Technology|B\.?S\.?|B\.?A\.?|B\.?E\.?|B\.?Tech))",
            r"(Master[\s]*(?:of|in|Science|Arts|Engineering|Technology|M\.?S\.?|M\.?A\.?|M\.?E\.?|M\.?Tech|MBA))",
            r"(Ph\.?D\.?|Doctor(?:ate)?[\s]*(?:of|in)?)",
            r"(Associate[\s]*(?:of|in|Degree|A\.?S\.?|A\.?A\.?))",
            r"(Diploma[\s]*(?:in|of)?)",
        ]
        year_pattern = r"(?:19|20)\d{2}\s*[-–]\s*(?:19|20)\d{2}|(?:19|20)\d{2}|Present|Current"
        for line in edu_text.split("\n"):
            line = line.strip()
            if not line:
                continue
            degree_found = None
            for pattern in degree_patterns:
                match = re.search(pattern, line, re.IGNORECASE)
                if match:
                    degree_found = match.group(1).strip()
                    break
            year_match = re.search(year_pattern, line, re.IGNORECASE)
            education.append({
                "degree": degree_found or line,
                "institution": line,
                "year": year_match.group(0) if year_match else None,
                "raw": line,
            })
    return education


def extract_experience(text: str, sections: Dict[str, Tuple[int, int]]) -> List[Dict[str, Any]]:
    experience = []
    lines = text.split("\n")
    if "experience" in sections:
        start, end = sections["experience"]
        exp_text = "\n".join(lines[start:end])
        year_pattern = r"(?:19|20)\d{2}\s*[-–]\s*(?:19|20)\d{2}|(?:19|20)\d{2}\s*[-–]\s*(?:Present|Current)|\d+\+?\s*(?:years?|yrs?)"
        current_entry: Dict[str, Any] = {"title": None, "company": None, "period": None, "description": []}
        for line in exp_text.split("\n"):
            line = line.strip()
            if not line:
                continue
            year_match = re.search(year_pattern, line, re.IGNORECASE)
            if year_match:
                if current_entry.get("title") or current_entry.get("description"):
                    experience.append(current_entry)
                current_entry = {
                    "title": line.replace(year_match.group(0), "").strip().strip("|-–—").strip(),
                    "company": None,
                    "period": year_match.group(0),
                    "description": [],
                }
            else:
                if current_entry.get("title") and not current_entry.get("company") and len(line) < 100:
                    current_entry["company"] = line
                elif len(line) > 5:
                    current_entry["description"].append(line)
        if current_entry.get("title") or current_entry.get("description"):
            experience.append(current_entry)
    return experience


def extract_projects(text: str, sections: Dict[str, Tuple[int, int]]) -> List[Dict[str, Any]]:
    projects = []
    lines = text.split("\n")
    if "projects" in sections:
        start, end = sections["projects"]
        proj_text = "\n".join(lines[start:end])
        current_project: Dict[str, Any] = {"name": None, "description": []}
        for line in proj_text.split("\n"):
            line = line.strip()
            if not line:
                continue
            if line.endswith(":") or (len(line) < 80 and not line.startswith(("•", "-", "*", "·"))):
                if current_project.get("name") or current_project.get("description"):
                    projects.append(current_project)
                current_project = {"name": line.rstrip(":"), "description": []}
            else:
                desc = re.sub(r"^[•\-\*·]\s*", "", line)
                if desc:
                    current_project["description"].append(desc)
        if current_project.get("name") or current_project.get("description"):
            projects.append(current_project)
    return projects


def extract_certifications(text: str, sections: Dict[str, Tuple[int, int]]) -> List[Dict[str, Any]]:
    certifications = []
    lines = text.split("\n")
    if "certifications" in sections:
        start, end = sections["certifications"]
        cert_text = "\n".join(lines[start:end])
        for line in cert_text.split("\n"):
            line = line.strip()
            if line and len(line) > 3:
                cert = re.sub(r"^[•\-\*·]\s*", "", line)
                if cert:
                    certifications.append({"name": cert, "raw": cert})
    return certifications


def extract_languages(text: str, sections: Dict[str, Tuple[int, int]]) -> List[str]:
    languages = []
    lines = text.split("\n")
    if "languages" in sections:
        start, end = sections["languages"]
        lang_text = "\n".join(lines[start:end])
        for line in lang_text.split("\n"):
            line = re.sub(r"^[•\-\*·]\s*", "", line.strip())
            if line and len(line) < 50:
                lang = line.split(":")[0].split("-")[0].split("(")[0].strip()
                if lang:
                    languages.append(lang)
    return languages


def parse_resume(file_path: str) -> Dict[str, Any]:
    raw_text, page_count = extract_text(file_path)
    cleaned_text = clean_text(raw_text)
    sections = find_section_boundaries(cleaned_text)
    email = extract_email(cleaned_text)
    phone = extract_phone(cleaned_text)
    linkedin = extract_linkedin(cleaned_text)
    github = extract_github(cleaned_text)
    portfolio = extract_portfolio_urls(cleaned_text, linkedin, github)
    name = extract_name_from_text(cleaned_text)
    skills = extract_skills_from_text(cleaned_text)
    education = extract_education(cleaned_text, sections)
    experience = extract_experience(cleaned_text, sections)
    projects = extract_projects(cleaned_text, sections)
    certifications = extract_certifications(cleaned_text, sections)
    languages = extract_languages(cleaned_text, sections)
    summary = None
    if "summary" in sections:
        start, end = sections["summary"]
        summary_lines = cleaned_text.split("\n")[start:end]
        summary = " ".join(line.strip() for line in summary_lines if line.strip() and line.strip().lower() not in RESUME_SECTIONS["summary"])

    return {
        "raw_text": raw_text,
        "cleaned_text": cleaned_text,
        "page_count": page_count,
        "parsed_data": {
            "name": name,
            "email": email,
            "phone": phone,
            "linkedin_url": linkedin,
            "github_url": github,
            "portfolio_url": portfolio,
            "summary": summary,
            "skills": skills,
            "education": education,
            "experience": experience,
            "projects": projects,
            "certifications": certifications,
            "languages": languages,
        },
    }
